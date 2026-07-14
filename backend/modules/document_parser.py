# backend/modules/document_parser.py
import os
import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
import openpyxl

from backend.core.config import settings

logger = logging.getLogger(__name__)


# ==================== 数据结构定义 ====================

@dataclass
class ParsedBlock:
    """
    解析后的原子文本块。
    保留结构元数据，供 text_splitter 进行语义切分时参考（如：同属一个标题下的段落不强行切断）。
    """
    text: str
    block_type: str = "paragraph"  # paragraph, title, table, list
    page_number: int = 0
    header_path: str = ""  # 面包屑导航，如 "第一章 > 1.1 引言"
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """文档解析最终产出物"""
    doc_id: str
    filename: str
    full_text: str  # 清洗后的纯文本（用于全局统计或降级兜底）
    blocks: List[ParsedBlock]  # 结构化文本块列表（核心资产）
    total_pages: int = 0
    total_chars: int = 0
    parse_status: str = "success"  # success, partial, failed
    error_msg: Optional[str] = None


# ==================== 解析器核心类 ====================

class DocumentParser:
    """
    企业级文档解析器。
    支持 PDF, DOCX, MD, TXT。具备防 OOM、防恶意文件、底层异常隔离能力。
    """

    # 防御性限制：防止恶意上传超大文件导致 Worker OOM 或 CPU 耗尽
    MAX_FILE_SIZE_MB = 50
    MAX_PDF_PAGES = 1000
    MAX_TXT_CHARS = 2_000_000  # 约 200 万字

    def parse(self, file_path: str, doc_id: str) -> ParsedDocument:
        """主入口：校验文件并路由到具体格式的解析器"""
        path = Path(file_path)
        filename = path.name

        # 1. 基础安全与存在性校验
        if not path.exists() or not path.is_file():
            return self._build_failed(doc_id, filename, "文件不存在或不可读")

        file_size_mb = path.stat().st_size / (1024 * 1024)
        if file_size_mb > self.MAX_FILE_SIZE_MB:
            return self._build_failed(doc_id, filename,
                                      f"文件过大 ({file_size_mb:.1f}MB)，上限为 {self.MAX_FILE_SIZE_MB}MB")

        # 2. 路由到具体解析策略
        ext = path.suffix.lower()
        try:
            if ext == ".pdf":
                return self._parse_pdf(file_path, doc_id, filename)
            elif ext == ".docx":
                return self._parse_docx(file_path, doc_id, filename)
            elif ext == ".doc":
                return self._parse_doc(file_path, doc_id, filename)
            elif ext == ".xlsx" or ext == ".xls":
                return self._parse_xlsx(file_path, doc_id, filename)
            elif ext == ".md":
                return self._parse_markdown(file_path, doc_id, filename)
            elif ext == ".txt":
                return self._parse_text(file_path, doc_id, filename)
            else:
                return self._build_failed(doc_id, filename, f"不支持的文件格式: {ext}")

        except Exception as e:
            # 兜底捕获：防止底层 C/C++ 库（如 MuPDF）抛出未预期的段错误导致 Celery Worker 崩溃
            logger.error(f"[Parser] 解析文件时发生未预期异常: {filename} | {e}", exc_info=True)
            return self._build_failed(doc_id, filename, f"解析引擎内部错误: {str(e)}")

    # ==================== 具体格式解析策略 ====================

    def _parse_pdf(self, file_path: str, doc_id: str, filename: str) -> ParsedDocument:
        """PDF 解析：使用 PyMuPDF，提取文本并尝试识别基础标题结构"""
        try:
            doc = fitz.open(file_path)
        except Exception as e:
            return self._build_failed(doc_id, filename, f"PDF 文件损坏或加密: {e}")

        if doc.page_count > self.MAX_PDF_PAGES:
            doc.close()
            return self._build_failed(doc_id, filename, f"PDF 页数 ({doc.page_count}) 超过限制 ({self.MAX_PDF_PAGES})")

        blocks: List[ParsedBlock] = []
        full_text_parts = []

        try:
            for page_num in range(doc.page_count):
                page = doc[page_num]
                # 使用 get_text("dict") 获取包含字体大小、位置的详细块信息，用于启发式标题识别
                page_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

                for block in page_dict.get("blocks", []):
                    if block["type"] == 0:  # 文本块
                        block_text = ""
                        max_font_size = 0
                        is_bold = False

                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                block_text += span["text"]
                                max_font_size = max(max_font_size, span["size"])
                                if "bold" in span["font"].lower() or "black" in span["font"].lower():
                                    is_bold = True
                            block_text += "\n"

                        block_text = block_text.strip()
                        if not block_text:
                            continue

                        # 启发式规则：字体较大且加粗，长度较短，判定为标题 (Title)
                        # 注：企业级复杂场景建议后续替换为 Docling 或 LayoutLM 模型
                        block_type = "paragraph"
                        if max_font_size > 14 and is_bold and len(block_text) < 50:
                            block_type = "title"

                        blocks.append(ParsedBlock(
                            text=block_text,
                            block_type=block_type,
                            page_number=page_num + 1,
                            metadata={"font_size": max_font_size, "is_bold": is_bold}
                        ))
                        full_text_parts.append(block_text)

        finally:
            # 【关键】必须显式关闭，防止底层 C 库内存泄漏
            doc.close()

        return self._build_success(doc_id, filename, blocks, full_text_parts, doc.page_count)

    def _parse_docx(self, file_path: str, doc_id: str, filename: str) -> ParsedDocument:
        """Word 解析：使用 python-docx，精准提取段落与内置样式（标题层级）"""
        try:
            doc = DocxDocument(file_path)
        except PackageNotFoundError:
            return self._build_failed(doc_id, filename, "DOCX 文件损坏或不是有效的 Word 文档")
        except Exception as e:
            return self._build_failed(doc_id, filename, f"DOCX 解析失败: {e}")

        blocks: List[ParsedBlock] = []
        full_text_parts = []
        current_header_path = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name.lower() if para.style else ""

            # 识别 Word 内置标题样式 (Heading 1, Heading 2, 标题 1, 标题 2 等)
            if "heading" in style_name or "标题" in style_name:
                block_type = "title"
                # 简单提取层级数字，更新面包屑导航
                level_match = re.search(r'\d', style_name)
                level = int(level_match.group()) if level_match else 1
                current_header_path = self._update_header_path(current_header_path, text, level)
            else:
                block_type = "paragraph"

            blocks.append(ParsedBlock(
                text=text,
                block_type=block_type,
                header_path=current_header_path
            ))
            full_text_parts.append(text)

        return self._build_success(doc_id, filename, blocks, full_text_parts, total_pages=0)

    def _parse_doc(self, file_path: str, doc_id: str, filename: str) -> ParsedDocument:
        """旧版 .doc 格式解析：尝试用 python-docx 解析，失败则尝试 subprocess 调用 antiword 或 textract"""
        # 策略1: 先用 python-docx 尝试（部分 .doc 可被正确读取）
        try:
            doc = DocxDocument(file_path)
            blocks: List[ParsedBlock] = []
            full_text_parts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    blocks.append(ParsedBlock(text=text, block_type="paragraph"))
                    full_text_parts.append(text)
            if blocks:
                logger.info(f"[Parser] .doc 文件通过 python-docx 解析成功: {filename}")
                return self._build_success(doc_id, filename, blocks, full_text_parts, total_pages=0)
        except Exception:
            pass

        # 策略2: 尝试 subprocess 调用 antiword（Linux/macOS 常见工具）
        import subprocess
        try:
            result = subprocess.run(
                ["antiword", file_path],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout.strip()
                blocks = [ParsedBlock(text=text, block_type="paragraph")]
                return self._build_success(doc_id, filename, blocks, [text], total_pages=0)
        except FileNotFoundError:
            pass
        except Exception:
            pass

        # 策略3: 尝试 subprocess 调用 textract
        try:
            result = subprocess.run(
                ["textract", file_path],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout.strip()
                blocks = [ParsedBlock(text=text, block_type="paragraph")]
                return self._build_success(doc_id, filename, blocks, [text], total_pages=0)
        except FileNotFoundError:
            pass
        except Exception:
            pass

        return self._build_failed(
            doc_id, filename,
            "无法解析 .doc 文件。请转换为 .docx 格式后重新上传，或安装 antiword/textract 工具。"
        )

    def _parse_xlsx(self, file_path: str, doc_id: str, filename: str) -> ParsedDocument:
        """Excel 解析：openpyxl 处理 .xlsx，xlrd 处理 .xls，多重降级策略"""
        ext = Path(file_path).suffix.lower()
        blocks: List[ParsedBlock] = []
        full_text_parts: List[str] = []

        # 策略1: openpyxl（适用于 .xlsx，部分 .xls 也能读）
        try:
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        except Exception:
            wb = None

        if wb is not None:
            try:
                for sheet_name in wb.sheetnames:
                    try:
                        ws = wb[sheet_name]
                        rows = list(ws.iter_rows(values_only=True))
                    except Exception:
                        continue
                    if not rows:
                        continue

                    header_path = f"工作表: {sheet_name}"
                    row_texts: List[str] = []
                    for row in rows:
                        cells = [str(cell) if cell is not None else "" for cell in row]
                        row_text = " | ".join(cells)
                        if row_text.strip():
                            row_texts.append(row_text)

                    sheet_text = "\n".join(row_texts)
                    if sheet_text.strip():
                        # 添加 Markdown 标题前缀，使 text_splitter 能识别 Sheet 边界
                        marked_text = f"## 工作表: {sheet_name}\n{sheet_text}"
                        blocks.append(ParsedBlock(
                            text=marked_text,
                            block_type="table",
                            header_path=header_path,
                        ))
                        full_text_parts.append(marked_text)
            except Exception as e:
                logger.warning(f"[Parser] openpyxl 处理工作表时出错: {filename} | {e}")
            finally:
                try:
                    wb.close()
                except Exception:
                    pass

            if full_text_parts:
                return self._build_success(doc_id, filename, blocks, full_text_parts, total_pages=0)

        # 策略2: xlrd（处理旧版 .xls 格式）
        if ext == ".xls":
            try:
                import xlrd
            except ImportError:
                logger.warning(f"[Parser] xlrd 未安装，无法解析 .xls 文件: {filename}")
                return self._build_failed(doc_id, filename, "无法解析 .xls 文件，请安装 xlrd 或转换为 .xlsx 格式后重新上传")

            try:
                wb_xl = xlrd.open_workbook(file_path)
            except Exception as e:
                return self._build_failed(doc_id, filename, f".xls 文件损坏或无法打开: {e}")

            try:
                for sheet_idx in range(wb_xl.nsheets):
                    ws = wb_xl.sheet_by_index(sheet_idx)
                    sheet_name = ws.name
                    header_path = f"工作表: {sheet_name}"
                    row_texts = []
                    for row_idx in range(ws.nrows):
                        cells = [str(ws.cell_value(row_idx, col)) if ws.cell_value(row_idx, col) != "" else "" for col in range(ws.ncols)]
                        row_text = " | ".join(cells)
                        if row_text.strip():
                            row_texts.append(row_text)

                    sheet_text = "\n".join(row_texts)
                    if sheet_text.strip():
                        # 添加 Markdown 标题前缀，使 text_splitter 能识别 Sheet 边界
                        marked_text = f"## 工作表: {sheet_name}\n{sheet_text}"
                        blocks.append(ParsedBlock(
                            text=marked_text,
                            block_type="table",
                            header_path=header_path,
                        ))
                        full_text_parts.append(marked_text)
            except Exception as e:
                logger.warning(f"[Parser] xlrd 处理工作表时出错: {filename} | {e}")
            finally:
                try:
                    wb_xl.release_resources()
                except Exception:
                    pass

            if full_text_parts:
                return self._build_success(doc_id, filename, blocks, full_text_parts, total_pages=0)

        # 策略3: pandas 兜底（最通用但较重）
        try:
            import pandas as pd
        except ImportError:
            pass
        else:
            try:
                if ext == ".xlsx" or ext == ".xls":
                    xl_file = pd.ExcelFile(file_path)
                    for sheet_name in xl_file.sheet_names:
                        df = pd.read_excel(xl_file, sheet_name=sheet_name, dtype=str)
                        if df.empty:
                            continue
                        header_path = f"工作表: {sheet_name}"
                        # 包含表头
                        rows_with_header = [df.columns.tolist()] + df.values.tolist()
                        row_texts = [" | ".join([str(c) if str(c) != "nan" else "" for c in row]) for row in rows_with_header]
                        sheet_text = "\n".join(row_texts)
                        if sheet_text.strip():
                            # 添加 Markdown 标题前缀
                            marked_text = f"## 工作表: {sheet_name}\n{sheet_text}"
                            blocks.append(ParsedBlock(
                                text=marked_text,
                                block_type="table",
                                header_path=header_path,
                            ))
                            full_text_parts.append(marked_text)
                    if full_text_parts:
                        return self._build_success(doc_id, filename, blocks, full_text_parts, total_pages=0)
            except Exception as e:
                logger.warning(f"[Parser] pandas 处理 Excel 出错: {filename} | {e}")

        return self._build_failed(doc_id, filename, "Excel 文件无法解析，请确认文件未损坏")

    def _parse_markdown(self, file_path: str, doc_id: str, filename: str) -> ParsedDocument:
        """Markdown 解析：基于正则按行解析标题层级与段落"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            return self._build_failed(doc_id, filename, "Markdown 文件编码非 UTF-8，无法解析")

        if len(content) > self.MAX_TXT_CHARS:
            return self._build_failed(doc_id, filename, "文本内容过长")

        blocks: List[ParsedBlock] = []
        full_text_parts = []
        current_header_path = ""
        current_paragraph = []

        def flush_paragraph():
            if current_paragraph:
                text = "\n".join(current_paragraph).strip()
                if text:
                    blocks.append(ParsedBlock(text=text, block_type="paragraph", header_path=current_header_path))
                    full_text_parts.append(text)
                current_paragraph.clear()

        for line in content.split("\n"):
            header_match = re.match(r'^(#{1,6})\s+(.*)', line)
            if header_match:
                flush_paragraph()
                level = len(header_match.group(1))
                title_text = header_match.group(2).strip()
                current_header_path = self._update_header_path(current_header_path, title_text, level)
                blocks.append(ParsedBlock(text=title_text, block_type="title", header_path=current_header_path))
                full_text_parts.append(title_text)
            else:
                if not line.strip():
                    flush_paragraph()
                else:
                    current_paragraph.append(line)

        flush_paragraph()
        return self._build_success(doc_id, filename, blocks, full_text_parts, total_pages=0)

    def _parse_text(self, file_path: str, doc_id: str, filename: str) -> ParsedDocument:
        """纯文本解析：直接读取，作为一个巨大的段落或按双换行符切分"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            return self._build_failed(doc_id, filename, "TXT 文件编码非 UTF-8")

        if len(content) > self.MAX_TXT_CHARS:
            return self._build_failed(doc_id, filename, "文本内容过长")

        blocks = []
        # 按连续两个换行符切分段落
        paragraphs = re.split(r'\n\s*\n', content)
        for p in paragraphs:
            text = p.strip()
            if text:
                blocks.append(ParsedBlock(text=text, block_type="paragraph"))

        return self._build_success(doc_id, filename, blocks, [content], total_pages=0)

    # ==================== 内部辅助方法 ====================

    def _update_header_path(self, current_path: str, new_title: str, level: int) -> str:
        """维护标题面包屑路径，例如：'第一章 > 1.1 背景'"""
        parts = current_path.split(" > ") if current_path else []
        # 截断到当前层级的前一级
        parts = parts[:level - 1]
        parts.append(new_title)
        return " > ".join(parts)

    def _build_success(self, doc_id: str, filename: str, blocks: List[ParsedBlock],
                       text_parts: List[str], total_pages: int) -> ParsedDocument:
        full_text = "\n\n".join(text_parts)
        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            full_text=full_text,
            blocks=blocks,
            total_pages=total_pages,
            total_chars=len(full_text),
            parse_status="success"
        )

    def _build_failed(self, doc_id: str, filename: str, error_msg: str) -> ParsedDocument:
        logger.warning(f"[Parser] 解析失败: {filename} | 原因: {error_msg}")
        return ParsedDocument(
            doc_id=doc_id,
            filename=filename,
            full_text="",
            blocks=[],
            parse_status="failed",
            error_msg=error_msg
        )


# 实例化单例供外部调用
document_parser = DocumentParser()